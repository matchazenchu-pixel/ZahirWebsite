from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = r"C:\Users\Ilyas\Documents\Website Zahir\Rancangan_Deployment_Aplikasi_Website_Zahir.docx"

NAVY = "173B5F"
BLUE = "DCEAF7"
LIGHT = "F5F8FB"
GRAY = "5B6573"
BORDER = "B7C7D9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = "w:" + edge
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key, value in edge_data.items():
                element.set(qn("w:" + key), str(value))


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + side))
        if node is None:
            node = OxmlElement("w:" + side)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def format_run(run, size=11, bold=False, color="000000", font="Arial"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_para(paragraph, before=0, after=0, line=1.3, align=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_text(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    format_run(run, **kwargs)
    return run


def add_body(doc, text):
    p = doc.add_paragraph()
    set_para(p, after=7, line=1.3, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.first_line_indent = Cm(0.75)
    add_text(p, text, size=11)
    return p


def add_heading(doc, number, title):
    p = doc.add_paragraph()
    set_para(p, before=12, after=6, line=1.3)
    add_text(p, f"{number}. {title}", size=13, bold=True, color=NAVY)
    p.paragraph_format.keep_with_next = True
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for i, label in enumerate(headers):
        cell = table.rows[0].cells[i]
        if widths:
            cell.width = Cm(widths[i])
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        set_cell_border(cell, top={"val": "single", "sz": "6", "color": NAVY}, bottom={"val": "single", "sz": "6", "color": NAVY}, left={"val": "single", "sz": "6", "color": NAVY}, right={"val": "single", "sz": "6", "color": NAVY})
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        set_para(p, line=1.3, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(p, label, size=10, bold=True, color="FFFFFF")
    for r, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cell = cells[i]
            if widths:
                cell.width = Cm(widths[i])
            if r % 2 == 1:
                set_cell_shading(cell, LIGHT)
            set_cell_margins(cell)
            set_cell_border(cell, top={"val": "single", "sz": "4", "color": BORDER}, bottom={"val": "single", "sz": "4", "color": BORDER}, left={"val": "single", "sz": "4", "color": BORDER}, right={"val": "single", "sz": "4", "color": BORDER})
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            set_para(p, line=1.3, align=WD_ALIGN_PARAGRAPH.LEFT)
            add_text(p, value, size=9.5)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    set_para(p, after=3, line=1.2)
    add_text(p, label, size=10.5, bold=True, color=NAVY)
    add_text(p, value, size=10.5)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
styles["Normal"].font.size = Pt(11)

# Header and footer
header = section.header.paragraphs[0]
set_para(header, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_text(header, "Rancangan Deployment Aplikasi | Website Zahir", size=8.5, color=GRAY)
header_border = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "4")
bottom.set(qn("w:space"), "4")
bottom.set(qn("w:color"), "B7C7D9")
header_border.append(bottom)
header._p.get_or_add_pPr().append(header_border)
footer = section.footer.paragraphs[0]
set_para(footer, before=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
add_text(footer, "Dokumen Rancangan — Deployment Aplikasi", size=8.5, color=GRAY)

# Title
p = doc.add_paragraph()
set_para(p, before=8, after=4, line=1.3, align=WD_ALIGN_PARAGRAPH.CENTER)
add_text(p, "RANCANGAN DEPLOYMENT APLIKASI", size=18, bold=True, color=NAVY)
p = doc.add_paragraph()
set_para(p, after=12, line=1.3, align=WD_ALIGN_PARAGRAPH.CENTER)
add_text(p, "WEBSITE ZAHIR", size=15, bold=True, color=NAVY)

info = doc.add_table(rows=2, cols=2)
info.alignment = WD_TABLE_ALIGNMENT.CENTER
info.autofit = False
for row, values in zip(info.rows, [("Mata Kuliah yang Dikonversi", "Deployment Aplikasi"), ("Objek Magang", "Website Zahir")]):
    for i, value in enumerate(values):
        cell = row.cells[i]
        cell.width = Cm(6.5 if i == 0 else 9.5)
        set_cell_margins(cell, top=85, bottom=85)
        set_cell_border(cell, top={"val": "single", "sz": "4", "color": BORDER}, bottom={"val": "single", "sz": "4", "color": BORDER}, left={"val": "single", "sz": "4", "color": BORDER}, right={"val": "single", "sz": "4", "color": BORDER})
        if i == 0:
            set_cell_shading(cell, BLUE)
        p = cell.paragraphs[0]
        set_para(p, line=1.3)
        add_text(p, value, size=10, bold=(i == 0), color=NAVY if i == 0 else "000000")

add_heading(doc, "1", "Tujuan Rancangan")
add_body(doc, "Rancangan ini disusun untuk menghasilkan Website Zahir yang dapat diakses secara online melalui proses deployment. Website yang sebelumnya berjalan pada lingkungan lokal akan dipublikasikan menggunakan platform Vercel agar dapat digunakan, diuji, dan ditunjukkan kepada pihak terkait melalui tautan publik.")

add_heading(doc, "2", "Data Magang yang Diolah")
add_body(doc, "Data yang diperoleh dari tempat magang diolah menjadi konten dan aset pada Website Zahir. Data tersebut tidak berdiri sebagai arsip, melainkan menjadi materi yang ditampilkan dalam aplikasi web.")
add_table(doc, ["Jenis Data", "Pengolahan dalam Website"], [
    ("Profil usaha/organisasi", "Disusun menjadi informasi identitas dan gambaran Website Zahir."),
    ("Informasi layanan atau produk", "Dikelompokkan untuk ditampilkan pada bagian layanan/produk."),
    ("Konten teks", "Disunting dan disesuaikan menjadi isi halaman yang ringkas serta mudah dibaca."),
    ("Aset visual", "Dioptimalkan dan ditempatkan sebagai elemen pendukung antarmuka website."),
    ("Struktur halaman", "Diterapkan menjadi navigasi dan susunan halaman yang dapat diakses pengguna."),
], widths=[5.3, 10.7])

add_heading(doc, "3", "Bentuk Output yang Direncanakan")
add_body(doc, "Output utama adalah Website Zahir yang telah ter-deploy dan dapat diakses melalui internet. Produk akhir tidak hanya berupa source code, tetapi juga mencakup tautan aplikasi, repository proyek, serta dokumentasi langkah deployment dan pengujian akses.")
add_table(doc, ["Output", "Keterangan"], [
    ("Website online", "Website Zahir tersedia pada hosting Vercel dan dapat dibuka melalui URL publik."),
    ("Tautan deployment", "URL aplikasi digunakan sebagai bukti bahwa website telah dipublikasikan."),
    ("Repository proyek", "Source code disimpan pada repository Git untuk mendukung version control dan integrasi deployment."),
    ("Dokumentasi deployment", "Berisi catatan konfigurasi, tahapan publikasi, serta bukti proses deployment."),
    ("Dokumentasi pengujian", "Berisi hasil pengecekan akses, navigasi, tampilan, dan responsivitas setelah publikasi."),
], widths=[5.3, 10.7])

add_heading(doc, "4", "Alur Rancangan Deployment")
add_body(doc, "Deployment dilakukan secara bertahap agar versi yang dipublikasikan merupakan versi website yang siap diakses. Alur yang digunakan adalah sebagai berikut.")
add_table(doc, ["Tahap", "Kegiatan", "Hasil yang Diharapkan"], [
    ("1. Finalisasi aplikasi", "Memastikan struktur halaman, konten, aset, dan navigasi Website Zahir telah siap.", "Versi aplikasi siap dipublikasikan."),
    ("2. Pengelolaan source code", "Menyimpan source code proyek pada repository Git.", "Repository menjadi sumber deployment dan riwayat perubahan."),
    ("3. Integrasi hosting", "Menghubungkan repository proyek dengan akun Vercel.", "Vercel dapat membaca dan membangun aplikasi."),
    ("4. Konfigurasi deployment", "Menentukan pengaturan build, direktori publik, dan konfigurasi yang diperlukan oleh proyek.", "Proses build berjalan sesuai struktur aplikasi."),
    ("5. Publikasi aplikasi", "Menjalankan proses deployment melalui Vercel.", "Website memperoleh URL publik."),
    ("6. Pengujian pascadeployment", "Menguji URL, navigasi, tampilan, dan responsivitas pada perangkat/ukuran layar yang relevan.", "Website dapat digunakan dengan baik setelah online."),
    ("7. Dokumentasi", "Mencatat URL, konfigurasi penting, bukti proses, dan hasil pengujian.", "Tersedia bukti pelaksanaan deployment."),
], widths=[2.3, 8.0, 5.7])

add_heading(doc, "5", "Rencana Pengujian Setelah Deployment")
add_body(doc, "Pengujian dilakukan setelah aplikasi memperoleh URL publik. Hasil aktual pengujian dicatat pada dokumentasi deployment setelah proses publikasi dilakukan.")
add_table(doc, ["Aspek yang Diuji", "Kriteria Keberhasilan", "Bukti yang Disiapkan"], [
    ("Akses website", "URL publik dapat dibuka tanpa menjalankan aplikasi secara lokal.", "Tangkapan layar halaman website dan URL deployment."),
    ("Navigasi", "Tautan menu dan tombol mengarahkan pengguna ke halaman atau bagian yang sesuai.", "Tangkapan layar tiap alur navigasi utama."),
    ("Tampilan", "Konten, gambar, dan komponen halaman tampil tanpa kerusakan visual yang mengganggu.", "Tangkapan layar halaman utama dan halaman pendukung."),
    ("Responsivitas", "Tampilan tetap mudah digunakan pada layar desktop dan perangkat mobile.", "Tangkapan layar pada ukuran layar yang berbeda."),
], widths=[4.0, 7.0, 5.0])

add_heading(doc, "6", "Indikator Keberhasilan")
add_table(doc, ["Indikator", "Target"], [
    ("Aplikasi terpublikasi", "Website Zahir berhasil dideploy pada Vercel."),
    ("Akses publik", "Tersedia URL yang dapat diakses melalui browser."),
    ("Kesiapan fungsi dasar", "Halaman dan navigasi utama dapat digunakan setelah deployment."),
    ("Dokumentasi", "Tersedia catatan deployment serta bukti pengujian aplikasi."),
], widths=[5.3, 10.7])

add_heading(doc, "7", "Gambaran Akhir")
add_body(doc, "Melalui rancangan ini, mata kuliah Deployment Aplikasi dikonversi melalui kegiatan publikasi Website Zahir dari lingkungan pengembangan ke layanan hosting online. Hasil akhirnya berupa aplikasi website yang dapat diakses publik, disertai tautan deployment dan dokumentasi proses sebagai bukti pelaksanaan kegiatan magang.")

# Core document metadata
doc.core_properties.title = "Rancangan Deployment Aplikasi Website Zahir"
doc.core_properties.subject = "Konversi mata kuliah Deployment Aplikasi"
doc.core_properties.author = "【Nama Mahasiswa】"

doc.save(OUTPUT)
print(OUTPUT)
